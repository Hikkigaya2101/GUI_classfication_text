from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.Qt import *
from PyQt5.QtWidgets import QMainWindow, QApplication
from PyQt5.QtGui import QIcon
import sys
import pymorphy2_dicts_ru
from selenium import webdriver
import time
import os
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
import pymorphy2
import re
from docx import Document
from docx.shared import Cm
from docx.shared import Pt
from docx.shared import Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import  string
import  shutil
from threading import *
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.firefox.options import Options
from bs4 import BeautifulSoup as BS
import requests
import random
import numpy as np

from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
from matplotlib.backends.backend_qt5agg import NavigationToolbar2QT as NavigationToolbar
from matplotlib.figure import Figure


from requests_html import HTMLSession
#LOGIN
class Ui_Form(object):
    def setupUi(self, Form):
        Form.setObjectName("Form")
        Form.resize(309, 430)
        Form.setStyleSheet("background-color:rgba(16,31,41,240);\n"
"border-radius:10px;")
        Form.setWindowIcon(QtGui.QIcon('icon.png'))
        self.label_2 = QtWidgets.QLabel(Form)
        self.label_2.setEnabled(True)
        self.label_2.setGeometry(QtCore.QRect(90, 50, 128, 128))
        self.label_2.setStyleSheet("")
        self.label_2.setObjectName("label_2")
        self.lineEdit = QtWidgets.QLineEdit(Form)
        self.lineEdit.setGeometry(QtCore.QRect(30, 240, 250, 30))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.lineEdit.setFont(font)
        self.lineEdit.setStyleSheet("background-color:rgba(0,0,0,0);\n"
"border:1px solid rgba(0,0,0,0);\n"
"border-bottom-color:rgba(46,82,101,255);\n"
"color:rgb(255,255,255,255);\n"
"padding-bottom:7px;")
        self.lineEdit.setObjectName("lineEdit")
        self.lineEdit_2 = QtWidgets.QLineEdit(Form)
        self.lineEdit_2.setGeometry(QtCore.QRect(30, 290, 250, 30))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.lineEdit_2.setFont(font)
        self.lineEdit_2.setStyleSheet("background-color:rgba(0,0,0,0);\n"
"border:1px solid rgba(0,0,0,0);\n"
"border-bottom-color:rgba(46,82,101,255);\n"
"color:rgb(255,255,255,255);\n"
"padding-bottom:7px;")
        self.lineEdit_2.setEchoMode(QtWidgets.QLineEdit.Password)
        self.lineEdit_2.setObjectName("lineEdit_2")
        self.logbutton = QtWidgets.QPushButton(Form)
        self.logbutton.setGeometry(QtCore.QRect(30, 350, 250, 40))
        font = QtGui.QFont()
        font.setPointSize(10)
        font.setBold(True)
        font.setWeight(75)
        self.logbutton.setFont(font)
        self.logbutton.setStyleSheet("QPushButton#pushButton{\n"
"background-color:rgba(2,65,118,255);\n"
"color:rgba(255,255,255,200);\n"
"border-radius:5px;\n"
"}\n"
"QPushButton#logbutton:pressed{\n"
"padding-left:5px;\n"
"padding-top:5px;\n"
"background-color:rgba(2,65,118,100);\n"
"background-position:calc(100%-10px)center;\n"
"}\n"
"QPushButton#logbutton{\n"
"background-color:rgba(2,65,118,200);\n"
"\n"
"}\n"
"")
        self.logbutton.setObjectName("logbutton")
        self.retranslateUi(Form)
        QtCore.QMetaObject.connectSlotsByName(Form)
#Получение текста из LINE
    def get_text(self):
        global login,pas
        login = self.lineEdit.text()
        pas = self.lineEdit_2.text()
        return login,pas
    def retranslateUi(self, Form):
        _translate = QtCore.QCoreApplication.translate
        Form.setWindowTitle(_translate("Form", "Авторизация"))
        self.label_2.setText(_translate("MainWindow", "<html><head/><body><p><img src=\"C:\\selenium\\logo.png\"/></p></body></html>"))
        self.lineEdit.setPlaceholderText(_translate("Form", "Логин"))
        self.lineEdit_2.setPlaceholderText(_translate("Form", "Пароль"))
        self.logbutton.setText(_translate("Form", "Войти"))
        self.logbutton.clicked.connect(lambda:self.get_text())
#MAIN_FORM
class Ui_MainWindow(object):
    def setupUi(self, MainWindow):
        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(870, 494)
        MainWindow.setStyleSheet("color:#fff;\n"
                                 "background-color: rgb(33,43,51);\n"
                                 "font-size:15px;")
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.horizontalLayout = QtWidgets.QHBoxLayout(self.centralwidget)
        self.horizontalLayout.setObjectName("horizontalLayout")
        self.frame = QtWidgets.QFrame(self.centralwidget)
        self.frame.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.frame.setFrameShadow(QtWidgets.QFrame.Raised)
        self.frame.setObjectName("frame")
        self.verticalLayout_2 = QtWidgets.QVBoxLayout(self.frame)
        self.verticalLayout_2.setObjectName("verticalLayout_2")
        self.groupBox = QtWidgets.QGroupBox(self.frame)
        self.groupBox.setObjectName("groupBox")
        self.verticalLayout = QtWidgets.QVBoxLayout(self.groupBox)
        self.verticalLayout.setObjectName("verticalLayout")
        self.radioButton_4 = QtWidgets.QRadioButton(self.groupBox)
        self.radioButton_4.setObjectName("radioButton_4")
        self.verticalLayout.addWidget(self.radioButton_4)
        self.radioButton = QtWidgets.QRadioButton(self.groupBox)
        self.radioButton.setObjectName("radioButton")
        self.verticalLayout.addWidget(self.radioButton)
        self.radioButton_2 = QtWidgets.QRadioButton(self.groupBox)
        self.radioButton_2.setObjectName("radioButton_2")
        self.verticalLayout.addWidget(self.radioButton_2)
        self.radioButton_3 = QtWidgets.QRadioButton(self.groupBox)
        self.radioButton_3.setObjectName("radioButton_3")
        self.verticalLayout.addWidget(self.radioButton_3)
        self.verticalLayout_2.addWidget(self.groupBox)
        self.dockbutton = QtWidgets.QPushButton(self.frame)
        self.dockbutton.setObjectName("dockbutton")
        self.verticalLayout_2.addWidget(self.dockbutton)
        self.classbutton = QtWidgets.QPushButton(self.frame)
        self.classbutton.setObjectName("classbutton")
        self.verticalLayout_2.addWidget(self.classbutton)
        self.parsbutton = QtWidgets.QPushButton(self.frame)
        self.parsbutton.setObjectName("parsbutton")
        self.verticalLayout_2.addWidget(self.parsbutton)
        self.horizontalLayout.addWidget(self.frame)

        self.widget = QtWidgets.QWidget(self.centralwidget)
        self.widget.setObjectName("widget")
        self.horizontalLayout.addWidget(self.widget)


        MainWindow.setCentralWidget(self.centralwidget)

        self.retranslateUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)
        #self.radioButton.toggled.connect(lambda: self.parser(self.radioButton))
        self.parsbutton.clicked.connect(lambda: self.start_parser())
        #self.classbutton.clicked.connect(lambda: self.Sheet._plot)
        self.dockbutton.clicked.connect(lambda: self.IAM())

    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "Selenium"))
        MainWindow.setWindowIcon(QtGui.QIcon('icon.png'))
        self.groupBox.setTitle(_translate("MainWindow", "Новостные порталы"))
        self.radioButton_4.setText(_translate("MainWindow", "rbk_rostov"))
        self.radioButton.setText(_translate("MainWindow", "161"))
        self.radioButton_2.setText(_translate("MainWindow", "ro_today"))
        self.radioButton_3.setText(_translate("MainWindow", "rnd"))
        self.dockbutton.setText(_translate("MainWindow", "Документ"))
        self.classbutton.setText(_translate("MainWindow", "Классификация"))
        self.parsbutton.setText(_translate("MainWindow", "Сбор"))
        #self.plainTextEdit.setPlainText(_translate("MainWindow", "СООБЩЕНИЯ"""))

    def start_parser(self):
        if self.radioButton_4.isChecked():
            t = Thread(target=self.parser)
            t.start()
        if self.radioButton_3.isChecked():
            t = Thread(target=self.parser_rnd)
            t.start()
        if self.radioButton_2.isChecked():
            t = Thread(target=self.parser_ro_today)
            t.start()
        if self.radioButton.isChecked():
            t = Thread(target=self.parser_rbk)
            t.start()
    def IAM(self):
        document = Document()
        style = document.styles['Normal']
        style.font.name = 'Times New Roman'
        document.add_heading('2.Деятельность руководителя субъекта РФ', 7)
        document.add_paragraph(' Основные результаты деятельности', style='List Number')
        document.add_paragraph('Общественное мнение о деятельности руководства субъекта РФ', style='List Number')
        document.add_paragraph('Промышленное производство', style='List Number')
        document.add_paragraph('Агропромышленный комплекс', style='List Number')
        document.add_paragraph('Инвестиции ', style='List Number')
        document.add_paragraph('Бюджет', style='List Number')
        document.add_paragraph('Уровень жизни населения', style='List Number')

        ###########
        res = os.listdir(r'/class/1.result')
        result = "C:\\selenium\\class\\1.result\\" + res[0]
        text_t = open(result, mode='r', encoding='utf-8')
        text_t = text_t.readlines()
        ###########
        document.add_heading('2.1. Основные результаты деятельности', 7)
        style.font.size = Pt(14)
        style.font.name = 'Times New Roman'
        p = document.add_paragraph(text_t)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        fmt = p.paragraph_format
        fmt.first_line_indent = Mm(15)
        fmt.space_after = Mm(10)
        ###########
        social = os.listdir(r'/class/2.social')
        social = "C:\\selenium\\class\\2.social\\" + social[0]
        text_s = open(social, mode='r', encoding='utf-8')
        text_s = text_s.readlines()

        ###########
        document.add_heading('2.2.Общественное мнение о деятельности руководства субъекта РФ', 7)
        p1 = document.add_paragraph(text_s)
        p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        fmt1 = p1.paragraph_format
        fmt1.first_line_indent = Mm(15)
        fmt1.space_after = Mm(10)
        ###########
        industrial = os.listdir(r'/class/3.industrial')
        industrial = "C:\\selenium\\class\\3.industrial\\" + industrial[0]
        text_i = open(industrial, mode='r', encoding='utf-8')
        text_i = text_i.readlines()

        ###########
        document.add_heading('2.3. Промышленное производство', 7)
        p1 = document.add_paragraph(text_i)
        p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        fmt1.first_line_indent = Mm(15)
        fmt1.space_after = Mm(10)
        ###########
        agro = os.listdir(r'/class/4.agro')
        agro = "C:\\selenium\\class\\4.agro\\" + agro[0]
        text_a = open(agro, mode='r', encoding='utf-8')
        text_a = text_a.readlines()

        ###########
        document.add_heading('2.4.Агропромышленный комплекс', 7)
        p1 = document.add_paragraph(text_a)
        p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        fmt1.first_line_indent = Mm(15)
        fmt1.space_after = Mm(10)
        ###########
        invest = os.listdir(r'/class/5.invest')
        invest = "C:\\selenium\\class\\5.invest\\" + invest[0]
        text_in = open(invest, mode='r', encoding='utf-8')
        text_in = text_in.readlines()

        ###########
        document.add_heading('2.5. Инвестиции', 7)
        p1 = document.add_paragraph(text_in)
        p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        fmt1.first_line_indent = Mm(15)
        fmt1.space_after = Mm(10)
        ###########
        budget = os.listdir(r'/class/6.budget')
        budget = "C:\\selenium\\class\\6.budget\\" + budget[0]
        text_b = open(budget, mode='r', encoding='utf-8')
        text_b = text_b.readlines()

        ###########
        document.add_heading('2.6. Бюджет', 7)
        p1 = document.add_paragraph(text_b)
        p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        fmt1.first_line_indent = Mm(15)
        fmt1.space_after = Mm(10)
        ###########
        level = os.listdir(r'/class/7.level')
        level = "C:\\selenium\\class\\7.level\\" + level[0]
        text_l = open(level, mode='r', encoding='utf-8')
        text_l = text_l.readlines()

        ###########
        document.add_heading('2.7. Уровень жизни населения', 7)
        p1 = document.add_paragraph(text_l)
        p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        fmt1.first_line_indent = Mm(15)
        fmt1.space_after = Mm(10)
        document.save(r'C:\selenium\IAM.docx')
        self.plainTextEdit.insertPlainText("Формирование документа завершено\n")
    def parser_rbk(self):
        BASE_DIR = r"C:/selenium/news"
        #proxies = { 'http': 'http://10.0.0.1:3128','https': 'http://10.0.0.1:3128',}
        r = requests.get('https://www.1rnd.ru/news')#, proxies=proxies)
        soup = BS(r.content, 'html.parser')
        a = soup.find_all("a", class_="c-news-block__title")
        #print(len(a))
        # i = len(a)
        text = []
        href = []
        for e in a:
            text.append(e.getText())
            href.append(e.attrs.get("href"))
        #print(href)
        news = []
        for h in href:
            r = requests.get(h)#, proxies=proxies)
            soup = BS(r.content, 'html.parser')
            a = soup.find("div", class_="col-12 col-lg-9")
            te = a.text
            news.append(str(te))
        for mas in range(1, 20):
            # print(mas)
            path = rf"{BASE_DIR}\{'_'.join(text[mas].split(' '))[:-4:]}.txt"
            #  path = pathlib.Path('C:/') / 'diplom' / mas + '.txt'
            with open(path, 'w', encoding="utf-8") as file:
                file.write(news[mas])
    def parser_ro_today(self):
        smi_header = """\n==/СМИ\n"""
        BASE_DIR = r"C:/selenium/news"
        #proxies = {'http': 'http://10.0.0.1:3128', 'https': 'http://10.0.0.1:3128',}
        r = requests.get('https://www.altairegion22.ru/')#, proxies=proxies)
        soup = BS(r.content, 'html.parser')
        today = soup.find("a", class_="today")
        href_today = today.attrs.get("href")
        region = 'https://www.altairegion22.ru' + href_today
        #print(region)
        new = requests.get(region)#, proxies=proxies)
        a = soup.find_all("a", class_="name")
        #print(len(a))
        # i = len(a)
        text = []
        href = []
        for e in a:
            text.append(e.getText())
            href.append(e.attrs.get("href"))
        #print(href)
        news = []
        for h in href:
            r = requests.get('https://www.altairegion22.ru' + h)#, proxies=proxies)
            soup = BS(r.content, 'html.parser')
            a = soup.find("div", class_="news_det_text")
            te = a.text
            news.append(str(te))
        for mas in range(1, 20):
            path = rf"{BASE_DIR}\{'_'.join(text[mas].split(' '))[:-4:]}.txt"
            #  path = pathlib.Path('C:/') / 'diplom' / mas + '.txt'
            with open(path, 'w', encoding="utf-8") as file:
                file.write(news[mas])
    def parser_rnd(self):
        options = webdriver.ChromeOptions()
        options.add_argument(
            "user-agent = Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/99.0.4844.51 Safari/537.36")
        options.add_argument("--disable-blink-features=AutomationControlled")
        driver = webdriver.Chrome(
            executable_path=r'C:/selenium/chromedriver/chromedriver.exe',
            options=options
        )
        try:
            driver.get("https://www.1rnd.ru/news")
            items = driver.find_elements(by=By.CLASS_NAME, value="c-news-block__image")
            items[0].click()
            pars1 = driver.find_element(by=By.CLASS_NAME, value="article-details__text")
            print(pars1.text)
            file = open("C:/selenium/news/1_rnd.txt", "w", encoding="utf-8")
            file.write(pars1.text)
            driver.back()
            items = driver.find_elements(by=By.CLASS_NAME, value="c-news-block__image")
            items[1].click()
            pars2 = driver.find_element(by=By.CLASS_NAME, value="article-details__text")
            file = open("C:/selenium/news/2_rnd.txt", "w", encoding="utf-8")
            file.write(pars2.text)
            driver.back()
            items = driver.find_elements(by=By.CLASS_NAME, value="c-news-block__image")
            items[2].click()
            pars3 = driver.find_element(by=By.CLASS_NAME, value="article-details__text")
            file = open("C:/selenium/news/3_rnd.txt", "w", encoding="utf-8")
            file.write(pars3.text)
            driver.back()
            items = driver.find_elements(by=By.CLASS_NAME, value="c-news-block__image")
            items[3].click()
            pars4 = driver.find_element(by=By.CLASS_NAME, value="article-details__text")
            file = open("C:/selenium/news/4_rnd.txt", "w", encoding='utf-8')
            file.write(pars4.text)
            driver.back()
            items = driver.find_elements(by=By.CLASS_NAME, value="c-news-block__image")
            items[4].click()
            pars5 = driver.find_element(by=By.CLASS_NAME, value="article-details__text")
            file = open("C:/selenium/news/5_rnd.txt", "w", encoding='utf-8')
            file.write(pars5.text)
            driver.back()
            items = driver.find_elements(by=By.CLASS_NAME, value="c-news-block__image")
            items[5].click()
            pars6 = driver.find_element(by=By.CLASS_NAME, value="article-details__text")
            file = open("C:/selenium/news/6_rnd.txt", "w", encoding='utf-8')
            file.write(pars6.text)
            driver.back()
            items = driver.find_elements(by=By.CLASS_NAME, value="c-news-block__image")
            items[6].click()
            pars7 = driver.find_element(by=By.CLASS_NAME, value="article-details__text")
            file = open("C:/selenium/news/7_rnd.txt", "w", encoding='utf-8')
            file.write(pars7.text)
            driver.back()
            items = driver.find_elements(by=By.CLASS_NAME, value="c-news-block__image")
            items[7].click()
            pars8 = driver.find_element(by=By.CLASS_NAME, value="article-details__text")
            file = open("C:/selenium/news/8_rnd.txt", "w", encoding='utf-8')
            file.write(pars8.text)
            driver.back()
            items = driver.find_elements(by=By.CLASS_NAME, value="c-news-block__image")
            items[8].click()
            pars9 = driver.find_element(by=By.CLASS_NAME, value="article-details__text")
            print(pars9.text)
            file = open("C:/selenium/news/9_rnd.txt", "w", encoding='utf-8')
            file.write(pars9.text)
            driver.back()
            items = driver.find_elements(by=By.CLASS_NAME, value="c-news-block__image")
            items[9].click()
            pars10 = driver.find_element(by=By.CLASS_NAME, value="article-details__text")
            print(pars10.text)
            file = open("C:/selenium/news/10.txt", "w", encoding='utf-8')
            file.write(pars10.text)
            driver.back()
        finally:
            driver.close()
            driver.quit()

    def parser(self):
        options = Options()
        options.binary_location = r'C:\Program Files\Mozilla Firefox\firefox.exe'
        from selenium.webdriver.firefox.firefox_binary import FirefoxBinary
        from selenium.webdriver.common.desired_capabilities import DesiredCapabilities
        firefox_capabilities = DesiredCapabilities.FIREFOX
        firefox_capabilities['marionette'] = True
        driver = webdriver.Firefox(options=options)
        try:
            driver.get("https://161.ru/text/")
            items = driver.find_elements(by=By.CLASS_NAME, value="h9Jmx")
            items[0].click()
            pars1 = driver.find_element(by=By.CLASS_NAME, value="Y4bXJ")
            print(pars1.text)
            file = open("C:\\selenium\\news\\1_161.txt", "w", encoding="utf-8")
            file.write(pars1.text)
            driver.back()
            items = driver.find_elements(by=By.CLASS_NAME, value="h9Jmx")
            items[1].click()
            pars2 = driver.find_element(by=By.CLASS_NAME, value="Y4bXJ")
            file = open("C:\\selenium\\news\\2_161.txt", "w", encoding="utf-8")
            file.write(pars2.text)
            driver.back()
            items = driver.find_elements(by=By.CLASS_NAME, value="h9Jmx")
            items[2].click()
            pars3 = driver.find_element(by=By.CLASS_NAME, value="Y4bXJ")
            file = open("C:\\selenium\\news\\3_161.txt", "w", encoding="utf-8")
            file.write(pars3.text)
            driver.back()
            items = driver.find_elements(by=By.CLASS_NAME, value="h9Jmx")
            items[3].click()
            pars4 = driver.find_element(by=By.CLASS_NAME, value="Y4bXJ")
            file = open("/news/4_rnd.txt", "w", encoding='utf-8')
            file.write(pars4.text)
            driver.back()
            items = driver.find_elements(by=By.CLASS_NAME, value="h9Jmx")
            items[4].click()
            pars5 = driver.find_element(by=By.CLASS_NAME, value="Y4bXJ")
            file = open("/news/5_rnd.txt", "w", encoding='utf-8')
            file.write(pars5.text)
            driver.back()
            items = driver.find_elements(by=By.CLASS_NAME, value="h9Jmx")
            items[5].click()
            pars6 = driver.find_element(by=By.CLASS_NAME, value="Y4bXJ")
            file = open("/news/6_rnd.txt", "w", encoding='utf-8')
            file.write(pars6.text)
            driver.back()
            items = driver.find_elements(by=By.CLASS_NAME, value="h9Jmx")
            items[6].click()
            pars7 = driver.find_element(by=By.CLASS_NAME, value="Y4bXJ")
            file = open("/news/7_rnd.txt", "w", encoding='utf-8')
            file.write(pars7.text)
            driver.back()
            items = driver.find_elements(by=By.CLASS_NAME, value="h9Jmx")
            items[7].click()
            pars8 = driver.find_element(by=By.CLASS_NAME, value="Y4bXJ")
            file = open("/news/8_rnd.txt", "w", encoding='utf-8')
            file.write(pars8.text)
            driver.back()
            items = driver.find_elements(by=By.CLASS_NAME, value="h9Jmx")
            items[8].click()
            pars9 = driver.find_element(by=By.CLASS_NAME, value="Y4bXJ")
            print(pars9.text)
            file = open("/news/9_rnd.txt", "w", encoding='utf-8')
            file.write(pars9.text)
            driver.back()
            items = driver.find_elements(by=By.CLASS_NAME, value="h9Jmx")
            items[9].click()
            pars10 = driver.find_element(by=By.CLASS_NAME, value="Y4bXJ")
            print(pars10.text)
            file = open("/news/10.txt", "w", encoding='utf-8')
            file.write(pars10.text)
            driver.back()
        finally:
            driver.close()
            driver.quit()
class MyMplCanvas(FigureCanvas):
    def __init__(self, *args, **kwargs):
        self.fig = Figure()
        super(MyMplCanvas, self).__init__(self.fig, *args, **kwargs)

    def plot(self, labels, men_means, women_means, x, width):  # !!!
        self.fig.clear()  # !!!

        self.ax = self.fig.add_subplot(111)  # !!!

        rects1 = self.ax.bar(x - width / 2, men_means, width, label='Аниме',color = '#1ab39c')
        rects2 = self.ax.bar(x + width / 2, women_means, width, label='Танки',color = '#ac1ab3d6')
        self.ax.set_ylabel('Количество сообщений',color = '#ffffff')
        self.ax.set_title('Классификация',color = '#ffffff')
        self.ax.set_xticks(x)
        self.ax.set_xticklabels(labels)
        self.ax.legend()
        self.ax.patch.set_facecolor('#212c32')
        self.fig.patch.set_facecolor('#212c32')
        self.draw()  # !!!


class Sheet(QMainWindow, Ui_MainWindow):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setupUi(self)
        self.classbutton.clicked.connect(lambda: self._plot())
        self.canavas = MyMplCanvas()
        self.toolbar = NavigationToolbar(self.canavas, self)

        self.layout = QVBoxLayout(self.widget)
        self.layout.addWidget(self.canavas)
        self.layout.addWidget(self.toolbar)

    def _plot(self):
        labels = ['G1', 'G2', 'G3', 'G4', 'G5']
        men_means = [random.randrange(1, 100) for _ in range(5)]
        women_means = [random.randrange(1, 100) for _ in range(5)]
        x = np.arange(len(labels))  # расположение столбиков
        width = 0.35  # толщина столбика
        print('reade')

        self.canavas.plot(labels, men_means, women_means, x, width)  # !!!






#INIT
app = QtWidgets.QApplication(sys.argv)
Dialog = QtWidgets.QWidget()
ui = Ui_Form()
ui.setupUi(Dialog)
Dialog.show()


def openOtherWindow():
    if login and pas == 'admin':
       global MainWindow
       MainWindow = QtWidgets.QMainWindow()
       ui = Ui_MainWindow()
       ui.setupUi(MainWindow)
       Dialog.close()
       #MainWindow.show()
       win = Sheet()
       win.show()
    else:print('error')
ui.logbutton.clicked.connect(lambda: openOtherWindow())
sys.exit(app.exec_())
